from flask import Flask, request, render_template, send_file
import pandas as pd
from docx import Document
import io
import matplotlib.pyplot as plt
from io import BytesIO
import docx.shared
from docx.oxml.ns import qn

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        department = request.form['department']
        year = request.form['year']
        section = request.form['section']
        exam_type = request.form['exam_type']
        file = request.files['marksheet']

        if file:
            df = pd.read_csv(file)
            doc = generate_report(df, department, year, section, exam_type)
            
            # Send the document as a file download
            return send_file(io.BytesIO(doc.getvalue()), as_attachment=True, download_name='result_analysis_report.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    return render_template('index.html')

def generate_report(df, department, year, section, exam_type):
    doc = Document()
    
    # Define font size
    font_size = docx.shared.Pt(14)
    
    # Add College Information and Overall Analysis on the First Page
    heading = doc.add_heading('Velammal Engineering College', 0)
    heading.runs[0].font.size = docx.shared.Pt(16)  # Larger font size for the heading

    p = doc.add_paragraph(f'Department: {department}')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Year of Study: {year}')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Section: {section}')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Exam Type: {exam_type}')
    p.runs[0].font.size = font_size

    # Define passing marks based on exam type
    if exam_type in ['CIE1', 'CIE2']:
        passing_marks = 25
        total_marks = 50
    elif exam_type == 'Model Examination':
        passing_marks = 50
        total_marks = 100
    else:
        raise ValueError('Invalid exam type')

    # Overall Analysis
    total_students = len(df)
    passed_all = len(df[(df['Java'] >= passing_marks) & (df['Python'] >= passing_marks) & (df['TOC'] >= passing_marks) & (df['DS'] >= passing_marks) & (df['CN'] >= passing_marks)])
    failed_all = total_students - passed_all
    overall_pass_percentage = (passed_all / total_students) * 100

    p = doc.add_paragraph(f'\nOverall Analysis:')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Total number of students: {total_students}')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Number of students passed in all subjects: {passed_all}')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Number of students did not pass in all subjects: {failed_all}')
    p.runs[0].font.size = font_size

    p = doc.add_paragraph(f'Overall pass percentage: {overall_pass_percentage:.2f}%')
    p.runs[0].font.size = font_size

    # Generate and add overall analysis chart
    fig, ax = plt.subplots()
    labels = ['Passed', 'Failed']
    sizes = [passed_all, failed_all]
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    
    # Save chart to a BytesIO object and add to document
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    plt.close(fig)
    
    img_stream.seek(0)
    doc.add_picture(img_stream, width=docx.shared.Inches(5.0))

    # Center the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = 1  # Center alignment
    
    # Add centered text below the image
    doc.add_paragraph('Overall Analysis Chart', style='Normal').alignment = 1

    # Add a Page Break
    doc.add_page_break()

    subjects = ['Java', 'Python', 'TOC', 'DS', 'CN']
    for subject in subjects:
        doc.add_heading(subject, level=1)
        doc.paragraphs[-1].runs[0].font.size = docx.shared.Pt(16)  # Larger font size for headings

        passed = df[df[subject] >= passing_marks]
        failed = df[df[subject] < passing_marks]

        num_passed = len(passed)
        num_failed = len(failed)
        pass_percentage = (num_passed / total_students) * 100

        p = doc.add_paragraph(f'Number of students passed: {num_passed}')
        p.runs[0].font.size = font_size

        p = doc.add_paragraph(f'Number of students failed: {num_failed}')
        p.runs[0].font.size = font_size

        p = doc.add_paragraph(f'Total pass percentage: {pass_percentage:.2f}%')
        p.runs[0].font.size = font_size

        # Generate and add subject-wise chart
        fig, ax = plt.subplots()
        labels = ['Passed', 'Failed']
        sizes = [num_passed, num_failed]
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
        ax.axis('equal')
        
        # Save chart to a BytesIO object and add to document
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close(fig)
        
        img_stream.seek(0)
        doc.add_picture(img_stream, width=docx.shared.Inches(5.0))

        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # Center alignment

        # Add centered text below the image
        doc.add_paragraph(f'{subject} Performance Chart', style='Normal').alignment = 1

        # Add a Page Break after each subject report
        doc.add_page_break()

    # Remove Overall Analysis on the Last Page
    # No additional content on the last page

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

if __name__ == '__main__':
    app.run(host='192.168.120.41', port=5000, debug=True)
